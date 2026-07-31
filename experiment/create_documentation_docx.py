"""Create an editable DOCX intended for import into Google Docs.

Requires python-docx and BeautifulSoup. The output uses genuine Word heading
styles, semantic tables and list styles so Google Docs can preserve the
document outline and editing structure.
"""

from __future__ import annotations

import argparse
from pathlib import Path

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from create_google_docs_import import render


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_SOURCE = ROOT / "DOKUMENTATION_ENTWURF.md"
DEFAULT_OUTPUT = ROOT / "DOKUMENTATION_GOOGLE_DOCS.docx"

NAVY = "123247"
BLUE = "1E607B"
TEAL = "16847F"
CYAN = "58B9BF"
SOFT = "EEF4F6"
LIGHT = "F2F6F7"
HAIR = "D6E0E5"
INK = RGBColor(0x17, 0x27, 0x33)
MUTED = RGBColor(0x60, 0x72, 0x7E)


def set_cell_fill(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), color)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_run(run, color: str) -> None:
    r_pr = run._r.get_or_add_rPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    r_pr.append(shading)


def paragraph_border(paragraph, *, left: str | None = None, bottom: str | None = None) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    if left:
        edge = OxmlElement("w:left")
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), "18")
        edge.set(qn("w:space"), "8")
        edge.set(qn("w:color"), left)
        borders.append(edge)
    if bottom:
        edge = OxmlElement("w:bottom")
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), "14")
        edge.set(qn("w:space"), "4")
        edge.set(qn("w:color"), bottom)
        borders.append(edge)


def paragraph_shading(paragraph, color: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    p_pr.append(shading)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def set_table_borders(table, color: str = HAIR, size: str = "5") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = OxmlElement(f"w:{edge_name}")
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), size)
        edge.set(qn("w:color"), color)
        borders.append(edge)


def set_repeat_header_text(cell) -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(8.5)


def configure_styles(document: Document) -> None:
    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.18
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")

    for name, size, color, before, after in (
        ("Title", 30, NAVY, 0, 16),
        ("Subtitle", 15, BLUE, 0, 22),
        ("Heading 1", 22, NAVY, 18, 15),
        ("Heading 2", 15, BLUE, 18, 7),
        ("Heading 3", 12, TEAL, 14, 5),
    ):
        style = styles[name]
        style.font.name = "Aptos Display" if name != "Subtitle" else "Aptos"
        style.font.size = Pt(size)
        style.font.bold = name != "Subtitle"
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style._element.rPr.rFonts.set(qn("w:eastAsia"), style.font.name)

    styles["Heading 1"].paragraph_format.page_break_before = True

    quote = styles["Quote"]
    quote.font.name = "Aptos"
    quote.font.size = Pt(10.5)
    quote.font.color.rgb = RGBColor(0x28, 0x41, 0x4F)
    quote.paragraph_format.left_indent = Cm(0.45)
    quote.paragraph_format.right_indent = Cm(0.2)
    quote.paragraph_format.space_before = Pt(8)
    quote.paragraph_format.space_after = Pt(10)

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(10.5)
        style.font.color.rgb = INK
        style.paragraph_format.space_after = Pt(4)


def configure_page(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    header = section.header.paragraphs[0]
    header.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = header.add_run("MULTISENSOR-LIDAR-OBJEKTERKENNUNG  ·  ENTWURF 0.1")
    run.font.name = "Aptos"
    run.font.size = Pt(7.5)
    run.font.bold = True
    run.font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    label = footer.add_run("DCAITI  ·  16. JULI 2026    |    ")
    label.font.name = "Aptos"
    label.font.size = Pt(7.5)
    label.font.color.rgb = MUTED
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)


def add_cover(document: Document) -> None:
    band = document.add_table(rows=1, cols=1)
    band.alignment = WD_TABLE_ALIGNMENT.CENTER
    band.autofit = False
    band.columns[0].width = Cm(17.2)
    cell = band.cell(0, 0)
    set_cell_fill(cell, TEAL)
    set_cell_margins(cell, top=130, bottom=130)
    p = cell.paragraphs[0]
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("TECHNISCHE UND WISSENSCHAFTLICHE DOKUMENTATION")
    run.font.name = "Aptos"
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    document.add_paragraph("")
    title = document.add_paragraph(style="Title")
    title.add_run("Multisensor-LiDAR-\nObjekterkennung in einer Tiefgarage")

    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.add_run(
        "Training, Cross-Validation und Bewertung von PointPillars und "
        "CenterPoint auf zwei Einzelsensoransichten und einer fusionierten "
        "Punktwolke"
    )

    summary = document.add_paragraph()
    paragraph_shading(summary, SOFT)
    paragraph_border(summary, left=CYAN)
    summary.paragraph_format.left_indent = Cm(0.35)
    summary.paragraph_format.right_indent = Cm(0.2)
    summary.paragraph_format.space_before = Pt(14)
    summary.paragraph_format.space_after = Pt(38)
    run = summary.add_run(
        "Experiment-held-out 3-Fold-Cross-Validation · Person, Fahrrad und "
        "Auto · Vergleich von os0, os1 und merged · Fehleranalyse und "
        "Laufzeiteinordnung"
    )
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x28, 0x41, 0x4F)

    meta = document.add_table(rows=2, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.autofit = True
    set_table_borders(meta, color=HAIR, size="4")
    values = (
        ("PROJEKT", "Multisensor LiDAR 3D Object Detection"),
        ("DOKUMENTSTATUS", "Erster Entwurf · Version 0.1"),
        ("STAND", "16. Juli 2026"),
        ("HAUPTMODELL", "PointPillars · KITTI-Finetuning"),
    )
    for cell, (label, value) in zip((c for row in meta.rows for c in row.cells), values):
        set_cell_fill(cell, "F5F8F9")
        set_cell_margins(cell, top=130, start=140, bottom=130, end=140)
        p = cell.paragraphs[0]
        r = p.add_run(label + "\n")
        r.font.size = Pt(7.5)
        r.font.bold = True
        r.font.color.rgb = RGBColor.from_string(TEAL)
        r = p.add_run(value)
        r.font.size = Pt(9.5)
        r.font.color.rgb = INK

    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_static_toc(document: Document, soup: BeautifulSoup) -> None:
    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(18)
    run = title.add_run("Inhaltsverzeichnis")
    run.font.name = "Aptos Display"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(NAVY)
    paragraph_border(title, bottom=TEAL)

    for heading in soup.find_all(["h1", "h2"]):
        level = 1 if heading.name == "h1" else 2
        p = document.add_paragraph()
        p.paragraph_format.left_indent = Cm(0 if level == 1 else 0.65)
        p.paragraph_format.space_after = Pt(5 if level == 1 else 3)
        run = p.add_run(heading.get_text(" ", strip=True))
        run.font.name = "Aptos"
        run.font.size = Pt(10.5 if level == 1 else 9)
        run.font.bold = level == 1
        run.font.color.rgb = RGBColor.from_string(NAVY if level == 1 else BLUE)


def add_inline(paragraph, node, *, bold=False, italic=False, code=False) -> None:
    if isinstance(node, NavigableString):
        text = str(node)
        if not text:
            return
        run = paragraph.add_run(text)
        run.bold = bold
        run.italic = italic
        if code:
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor.from_string(BLUE)
            shade_run(run, SOFT)
        return
    if not isinstance(node, Tag):
        return
    if node.name == "br":
        paragraph.add_run().add_break()
        return
    next_bold = bold or node.name in ("strong", "b")
    next_italic = italic or node.name in ("em", "i")
    next_code = code or node.name == "code"
    for child in node.children:
        add_inline(
            paragraph,
            child,
            bold=next_bold,
            italic=next_italic,
            code=next_code,
        )


def add_table(document: Document, html_table: Tag) -> None:
    rows = html_table.find_all("tr")
    if not rows:
        return
    column_count = max(len(row.find_all(["th", "td"], recursive=False)) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table)
    for row_index, html_row in enumerate(rows):
        cells = html_row.find_all(["th", "td"], recursive=False)
        for col_index, html_cell in enumerate(cells):
            cell = table.cell(row_index, col_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if row_index == 0:
                set_cell_fill(cell, NAVY)
            elif row_index % 2 == 0:
                set_cell_fill(cell, LIGHT)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            for child in html_cell.children:
                add_inline(paragraph, child)
            for run in paragraph.runs:
                run.font.name = "Aptos"
                run.font.size = Pt(8.5)
                if row_index == 0:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_repeat_table_header(table.rows[0])
    document.add_paragraph().paragraph_format.space_after = Pt(1)


def add_list(document: Document, html_list: Tag, ordered: bool) -> None:
    style = "List Number" if ordered else "List Bullet"
    for item in html_list.find_all("li", recursive=False):
        paragraph = document.add_paragraph(style=style)
        for child in item.children:
            if isinstance(child, Tag) and child.name in ("ul", "ol"):
                continue
            add_inline(paragraph, child)
        for nested in item.find_all(["ul", "ol"], recursive=False):
            add_list(document, nested, nested.name == "ol")


def add_content(document: Document, soup: BeautifulSoup) -> None:
    first_heading = True
    for node in soup.children:
        if isinstance(node, NavigableString):
            continue
        if not isinstance(node, Tag):
            continue
        if node.name in ("h1", "h2", "h3"):
            level = {"h1": 1, "h2": 2, "h3": 3}[node.name]
            paragraph = document.add_paragraph(style=f"Heading {level}")
            if first_heading and level == 1:
                paragraph.paragraph_format.page_break_before = True
                first_heading = False
            for child in node.children:
                add_inline(paragraph, child)
            if level == 1:
                paragraph_border(paragraph, bottom=TEAL)
        elif node.name == "p":
            paragraph = document.add_paragraph()
            for child in node.children:
                add_inline(paragraph, child)
        elif node.name == "blockquote":
            for html_p in node.find_all("p", recursive=False) or [node]:
                paragraph = document.add_paragraph(style="Quote")
                paragraph_shading(paragraph, SOFT)
                paragraph_border(paragraph, left=TEAL)
                for child in html_p.children:
                    add_inline(paragraph, child)
        elif node.name == "pre":
            paragraph = document.add_paragraph()
            paragraph_shading(paragraph, NAVY)
            paragraph_border(paragraph, left=CYAN)
            paragraph.paragraph_format.left_indent = Cm(0.25)
            paragraph.paragraph_format.right_indent = Cm(0.2)
            run = paragraph.add_run(node.get_text())
            run.font.name = "Consolas"
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        elif node.name == "table":
            add_table(document, node)
        elif node.name in ("ul", "ol"):
            add_list(document, node, node.name == "ol")
        elif node.name == "hr":
            paragraph = document.add_paragraph()
            paragraph_border(paragraph, bottom=HAIR)


def build_document(source_path: Path) -> Document:
    source = source_path.read_text(encoding="utf-8")
    body_html, _ = render(source)
    soup = BeautifulSoup(body_html, "html.parser")

    document = Document()
    configure_styles(document)
    configure_page(document)
    document.core_properties.title = (
        "Multisensor-LiDAR-Objekterkennung in einer Tiefgarage"
    )
    document.core_properties.subject = "Technische und wissenschaftliche Projektdokumentation"
    document.core_properties.keywords = "LiDAR, PointPillars, CenterPoint, Sensorfusion, Cross-Validation"
    document.core_properties.comments = "Google-Docs-kompatibler Importentwurf"

    add_cover(document)
    add_static_toc(document, soup)
    add_content(document, soup)
    return document


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()

    source = args.source.resolve()
    output = args.output.resolve()
    if not source.is_file():
        raise FileNotFoundError(source)
    output.parent.mkdir(parents=True, exist_ok=True)
    build_document(source).save(output)
    print(f"DOCX: {output} ({output.stat().st_size / 1024:.1f} KiB)")


if __name__ == "__main__":
    main()
